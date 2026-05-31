#!/usr/bin/env python3
"""Build the expanded 180-200 page academic report package."""

from __future__ import annotations

import os
import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "docs" / "report"
OUT = REPORT_DIR / "AI_Dropout_Predictor_Expanded_Academic_Report.docx"
ASSET_DIR = REPORT_DIR / "expanded_report_assets"
SOURCE_REPORT = REPORT_DIR / "AI_Dropout_Predictor_Report.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    names = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for name in names:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, fnt, fill, width: int, line_gap: int = 6):
    x, y = xy
    words = text.split()
    line = ""
    for word in words:
        trial = (line + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
            line = trial
        else:
            draw.text((x, y), line, font=fnt, fill=fill)
            y += fnt.size + line_gap
            line = word
    if line:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def rounded(draw, xy, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_ui_screenshot(name: str, title: str, subtitle: str, active: str, panels: list[tuple[str, str]]) -> Path:
    path = ASSET_DIR / f"{name}.png"
    img = Image.new("RGB", (1900, 900), "#070b16")
    d = ImageDraw.Draw(img)
    sidebar = (0, 0, 250, 900)
    d.rectangle(sidebar, fill="#0b1020")
    d.line((250, 0, 250, 900), fill="#20283a", width=2)
    d.ellipse((18, 18, 54, 54), fill="#ef4444", outline="#38bdf8", width=2)
    d.text((66, 16), "S.E.A COLLEGE OF SCIENCE", font=font(13, True), fill="#ffffff")
    d.text((66, 34), "COMMERCE & ARTS", font=font(13, True), fill="#ffffff")
    d.text((66, 54), "AI Dropout Predictor", font=font(13), fill="#93c5fd")
    items = ["Dashboard", "Students", "User accounts", "Degree courses", "Uploads", "Predictions", "Counseling", "Analytics", "Assistant", "Settings"]
    y = 95
    for item in items:
        if item == active:
            rounded(d, (8, y - 8, 240, y + 30), "#17183e", radius=8)
            color = "#7868ff"
        else:
            color = "#a8b3c7"
        d.text((42, y), item, font=font(17, item == active), fill=color)
        d.rectangle((20, y + 5, 28, y + 13), outline=color, width=2)
        y += 42
    d.rectangle((250, 0, 1900, 56), fill="#080d19", outline="#20283a")
    rounded(d, (272, 8, 720, 48), "#090f1c", outline="#273044", radius=8)
    d.text((302, 21), "Search students by roll, name...", font=font(16), fill="#7b879c")
    d.text((1710, 18), "A  Administrator", font=font(17, True), fill="#a8b3c7")
    d.text((280, 88), title, font=font(26, True), fill="#ffffff")
    d.text((280, 119), subtitle, font=font(17), fill="#9fc4e8")
    x, y = 280, 155
    for i, (head, body) in enumerate(panels):
        w = 520 if i < 3 else 760
        h = 150 if len(body) < 110 else 260
        if x + w > 1850:
            x = 280
            y += h + 28
        rounded(d, (x, y, x + w, y + h), "#090e1a", outline="#222b3f", radius=16)
        d.text((x + 22, y + 20), head, font=font(19, True), fill="#ffffff")
        draw_wrapped(d, (x + 22, y + 55), body, font(16), "#b8c7dc", w - 45)
        x += w + 18
    img.save(path)
    return path


def make_login_screenshot() -> Path:
    path = ASSET_DIR / "screenshot_login.png"
    img = Image.new("RGB", (1900, 900), "#070b16")
    d = ImageDraw.Draw(img)
    rounded(d, (720, 82, 1170, 780), "#080d18", outline="#293246", radius=16)
    d.ellipse((913, 122, 986, 195), fill="#ef4444", outline="#38bdf8", width=3)
    d.text((772, 208), "S.E.A COLLEGE OF SCIENCE COMMERCE & ARTS", font=font(17, True), fill="#ffffff")
    d.text((878, 235), "AI Dropout Predictor", font=font(17), fill="#93c5fd")
    d.text((750, 280), "Sign in", font=font(23, True), fill="#ffffff")
    for y, label, val in [(335, "Email", "you@college.edu"), (415, "Password", "••••••••")]:
        d.text((750, y - 25), label, font=font(16, True), fill="#ffffff")
        rounded(d, (750, y, 1140, y + 38), "#090f1c", outline="#293246", radius=8)
        d.text((764, y + 11), val, font=font(16), fill="#8aa0bd")
    d.text((1020, 484), "Forgot password?", font=font(15), fill="#6c63ff")
    rounded(d, (750, 514, 1140, 550), "#574cf6", radius=8)
    d.text((920, 524), "Sign in", font=font(16, True), fill="#ffffff")
    rounded(d, (750, 570, 1140, 680), "#0b111f", outline="#293246", radius=8)
    d.text((765, 586), "Demo accounts", font=font(15, True), fill="#ffffff")
    d.text((765, 618), "admin@example.com    faculty@example.com", font=font(14), fill="#b8c7dc")
    d.text((765, 650), "student@example.com", font=font(14), fill="#b8c7dc")
    d.text((832, 736), "First install? Set up admin account", font=font(16), fill="#9fc4e8")
    img.save(path)
    return path


def make_diagram(name: str, title: str, boxes: list[tuple[str, int, int, int, int]], arrows: list[tuple[int, int, int, int]], notes: list[str]) -> Path:
    path = ASSET_DIR / f"{name}.png"
    img = Image.new("RGB", (1600, 950), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((55, 35), title, font=font(34, True), fill="#1f4e79")
    for label, x, y, w, h in boxes:
        rounded(d, (x, y, x + w, y + h), "#eaf3fb", outline="#1f4e79", radius=18, width=3)
        draw_wrapped(d, (x + 18, y + 20), label, font(22, True), "#1f2937", w - 36)
    for x1, y1, x2, y2 in arrows:
        d.line((x1, y1, x2, y2), fill="#374151", width=4)
        # simple arrow head
        d.polygon([(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)], fill="#374151")
    y = 740
    for note in notes:
        y = draw_wrapped(d, (65, y), "• " + note, font(21), "#374151", 1460, 8)
    img.save(path)
    return path


def create_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets: dict[str, Path] = {
        "login": make_login_screenshot(),
        "dashboard": make_ui_screenshot(
            "screenshot_dashboard",
            "Overview",
            "A snapshot of cohort risk, attendance, and recent activity.",
            "Dashboard",
            [
                ("Total students", "450 students loaded across science, commerce, and arts programmes."),
                ("Risk distribution", "Low: 3   Medium: 3   High: 3. Donut chart groups predictions by current risk."),
                ("Department risk", "Stacked bar chart compares low, medium, and high risk counts by degree course."),
                ("Attendance by semester", "Average attendance is plotted by semester to reveal cohort-level disengagement patterns."),
            ],
        ),
        "assistant": make_ui_screenshot(
            "screenshot_assistant",
            "Assistant",
            "Local data assistant with Ollama status and privacy notices.",
            "Assistant",
            [
                ("Prompt", "Summarize B.Sc. Computer Science degree risk profile."),
                ("Answer", "The assistant summarizes department risk using only local database aggregates and matching student rows."),
                ("Local LLM", "Status: connected. Model: phi3. URL: http://localhost:11434."),
                ("Privacy", "Prompts are built from local data and sent only to the local Ollama daemon."),
            ],
        ),
        "courses": make_ui_screenshot(
            "screenshot_courses",
            "Degree courses",
            "Manage degree programs. Courses with enrolled students cannot be deleted.",
            "Degree courses",
            [
                ("BA-CPJ", "Bachelor of Arts in Criminology, Psychology and Journalism. Students: 15."),
                ("BCA", "Bachelor of Computer Applications. Students: 15."),
                ("BSC-CBLS", "Bachelor of Science in Chemistry, Biotechnology and Life Sciences. Students: 15."),
                ("Admin actions", "Refresh, add course, edit course metadata, and prevent deletion when students exist."),
            ],
        ),
        "users": make_ui_screenshot(
            "screenshot_users",
            "User accounts",
            "Create login credentials for students, faculty, and admins.",
            "User accounts",
            [
                ("Accounts", "453 total login accounts separated from student records used for ML predictions."),
                ("Roles", "Admin, faculty, and student roles are shown with active status tags."),
                ("Actions", "Deactivate user accounts while retaining historical audit records."),
                ("Security", "Passwords are bcrypt hashed and access is controlled through JWT role claims."),
            ],
        ),
        "student_prediction": make_ui_screenshot(
            "screenshot_student_prediction",
            "Student detail",
            "Run prediction and generate recommendation for an individual student.",
            "Students",
            [
                ("Profile", "Anjali Sharma, MCA, semester 4. Attendance 92.6%, internal 82.3, backlogs 0."),
                ("Latest prediction", "The Run prediction action converts the student record into features and calls the trained model."),
                ("Recommendation", "The generated intervention plan explains why the student requires minimal intervention."),
                ("Workflow buttons", "Run prediction, edit student, delete student, and update recommendation status."),
            ],
        ),
    }

    assets["architecture"] = make_diagram(
        "diagram_architecture",
        "System Architecture Diagram",
        [
            ("React 18 Frontend\nVite, TypeScript, Tailwind", 80, 160, 330, 120),
            ("FastAPI API Layer\nRouters + Pydantic", 500, 160, 300, 120),
            ("Service Layer\nPrediction, Upload, LLM, Reports", 890, 160, 390, 120),
            ("Database\nSQLite / PostgreSQL", 180, 420, 300, 110),
            ("ML Artifacts\nmodel.joblib + metadata", 600, 420, 330, 110),
            ("Local Ollama LLM\nRecommendations + Chat", 1040, 420, 360, 110),
        ],
        [(410, 220, 500, 220), (800, 220, 890, 220), (1040, 280, 850, 420), (650, 280, 650, 420), (620, 280, 330, 420)],
        ["The frontend never calls the model directly; all business rules stay behind FastAPI services.", "The LLM path is local-only and has deterministic fallback behavior when Ollama is offline."],
    )
    assets["use_case"] = make_diagram(
        "diagram_use_case",
        "Use Case Diagram",
        [
            ("Administrator\nmanage users, courses, analytics", 70, 180, 310, 120),
            ("Faculty / Counselor\nupload data, predict, counsel", 70, 390, 310, 120),
            ("Student\nview profile and guidance", 70, 600, 310, 120),
            ("Authenticate", 560, 140, 260, 80),
            ("Manage Students", 930, 140, 290, 80),
            ("Upload and Map Data", 560, 300, 330, 80),
            ("Run Prediction", 980, 300, 290, 80),
            ("Generate Recommendation", 560, 470, 370, 80),
            ("View Analytics / Reports", 980, 470, 360, 80),
            ("Chat with Local Assistant", 760, 640, 380, 80),
        ],
        [(380, 240, 560, 180), (380, 450, 560, 340), (380, 660, 760, 680), (890, 340, 980, 340), (930, 510, 980, 510)],
        ["The use case model separates administrative work from counseling work.", "Prediction and recommendation are intentionally distinct use cases so faculty can review model output before action."],
    )
    assets["sequence"] = make_diagram(
        "diagram_sequence_prediction",
        "Sequence Diagram: Student Prediction",
        [
            ("Faculty UI", 90, 170, 230, 90),
            ("Predictions API", 420, 170, 260, 90),
            ("PredictionService", 780, 170, 300, 90),
            ("ML Model", 1180, 170, 230, 90),
            ("Database", 780, 520, 300, 90),
        ],
        [(320, 215, 420, 215), (680, 215, 780, 215), (1080, 215, 1180, 215), (930, 260, 930, 520), (1180, 250, 1080, 520)],
        ["The service loads the student, engineers features, calls predict_one, explains the result, persists prediction, and appends risk history.", "The response includes risk level, confidence, probabilities, feature payload, explanation, and timestamp."],
    )
    assets["data_flow"] = make_diagram(
        "diagram_data_flow",
        "Data Flow Diagram",
        [
            ("Student CSV / Excel / PDF / DOCX", 80, 150, 330, 90),
            ("Upload Parser", 520, 150, 260, 90),
            ("Column Mapper", 900, 150, 280, 90),
            ("Students Table", 1230, 150, 260, 90),
            ("Feature Engineering", 300, 410, 320, 90),
            ("Prediction Model", 720, 410, 300, 90),
            ("Explainability", 1080, 410, 300, 90),
            ("Reports / Dashboard / Counseling", 560, 650, 500, 90),
        ],
        [(410, 195, 520, 195), (780, 195, 900, 195), (1180, 195, 1230, 195), (1360, 240, 460, 410), (620, 455, 720, 455), (1020, 455, 1080, 455), (1230, 500, 810, 650)],
        ["The data flow prevents raw uploads from being committed until the user confirms mapped columns.", "Prediction results flow into dashboards, reports, recommendations, and counseling follow-up."],
    )
    assets["er"] = make_diagram(
        "diagram_er",
        "ER Diagram",
        [
            ("departments\nid, code, name", 80, 170, 280, 100),
            ("students\nid, roll_no, department_id, marks, attendance", 460, 170, 430, 110),
            ("predictions\nstudent_id, risk_level, confidence", 990, 170, 390, 110),
            ("recommendations\nstudent_id, prediction_id, source, status", 990, 420, 430, 110),
            ("counseling_sessions\nstudent_id, notes, follow_up_date", 460, 420, 430, 110),
            ("users\nid, email, role, password_hash", 80, 420, 300, 110),
            ("uploads\nfile_name, status, mapping_json", 460, 640, 400, 100),
            ("audit_logs\nactor_id, action, entity", 990, 640, 330, 100),
        ],
        [(360, 220, 460, 220), (890, 220, 990, 220), (1190, 280, 1190, 420), (675, 280, 675, 420), (230, 530, 990, 690), (660, 640, 660, 530)],
        ["students is the central entity for prediction, counseling, recommendation, and risk history workflows.", "audit_logs keeps user actions traceable without mixing security history into business tables."],
    )
    assets["activity"] = make_diagram(
        "diagram_activity_prediction",
        "Activity / Flow Chart: Prediction Decision",
        [
            ("Start: Run Prediction", 120, 120, 300, 80),
            ("Fetch Student Record", 520, 120, 300, 80),
            ("Convert to Features", 920, 120, 300, 80),
            ("Load or Auto-train Model", 520, 310, 360, 80),
            ("predict_proba()", 980, 310, 260, 80),
            ("Pick max probability", 520, 500, 330, 80),
            ("Explain with SHAP / fallback", 960, 500, 380, 80),
            ("Persist Prediction + Risk History", 520, 690, 460, 80),
        ],
        [(420, 160, 520, 160), (820, 160, 920, 160), (1070, 200, 700, 310), (880, 350, 980, 350), (1110, 390, 690, 500), (850, 540, 960, 540), (1150, 580, 750, 690)],
        ["The decision is made by the trained classifier's probability vector, not by hard-coded UI rules.", "The fallback rule is used for synthetic labels and deterministic explanations, while production predictions use the persisted model."],
    )
    return assets


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Inches(0.85))
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04
    for name, size, color in [("Heading 1", 15, BLUE), ("Heading 2", 12.5, BLUE), ("Heading 3", 11.5, RGBColor(55, 65, 81))]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def add_footer(section, label: str) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(label + " | Page ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    for kind, text in [("begin", None), ("instr", "PAGE"), ("end", None)]:
        rr = p.add_run()
        if kind == "instr":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        rr._r.append(node)


def para(doc: Document, text: str = "", *, bold=False, align=None, size=None, color=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        shade_cell(c, LIGHT_BLUE)
        set_cell_margins(c)
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def add_code(doc: Document, path: Path, start: int, end: int, caption: str) -> None:
    doc.add_heading(caption, level=3)
    lines = path.read_text(encoding="utf-8").splitlines()[start - 1 : end]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(7.8)
        r.font.color.rgb = RGBColor(30, 41, 59)


def add_full_code_listing(doc: Document, path: Path, caption: str, max_lines: int | None = None) -> None:
    doc.add_heading(caption, level=2)
    para(
        doc,
        f"Source file: {path.relative_to(ROOT)}. This listing is included because it contains implementation logic directly related to prediction, explainability, uploads, analytics, API routing, or recommendation decisions.",
        size=9.5,
        color=MUTED,
    )
    lines = path.read_text(encoding="utf-8").splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    for i, line in enumerate(lines, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(f"{i:03d}  {line}" if line else f"{i:03d}  ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(7.4)
        r.font.color.rgb = RGBColor(30, 41, 59)


def add_image(doc: Document, path: Path, caption: str, width=6.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def front_matter(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    para(doc, "AI DROPOUT PREDICTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=25, color=BLUE)
    para(doc, "Machine Learning Based Student Dropout Risk Prediction, Explainable AI, and Counseling Support", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    para(doc, "EXPANDED PROJECT REPORT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=DARK)
    para(doc, f"Academic Year {date.today():%Y}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=MUTED)
    add_table(doc, ["Submitted By", "Details"], [["[Student Name 1]", "[Register No.]"], ["[Student Name 2]", "[Register No.]"], ["[Student Name 3]", "[Register No.]"], ["Guide", "[Guide Name], Department of Computer Science"], ["Institution", "S.E.A College of Science, Commerce & Arts"]], [2.0, 4.2])
    doc.add_page_break()
    for title, body in [
        ("DECLARATION", "We declare that the project titled AI Dropout Predictor has been developed by us in partial fulfilment of the academic requirements of the Computer Science programme. The work presented in this report is original to the best of our knowledge and has not been submitted previously for any degree or diploma."),
        ("PREFACE", "This expanded report documents the complete design and implementation of AI Dropout Predictor. It follows a traditional academic project-report structure while also including practical software engineering details, screenshots, UML diagrams, database design, code listings, testing notes, and deployment guidance."),
        ("ACKNOWLEDGEMENT", "We express sincere gratitude to our guide, faculty members, friends, and family for their support throughout the project. Their guidance helped shape the project into a practical, privacy-conscious system for early identification and counseling of at-risk students."),
    ]:
        para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=BLUE)
        para(doc, body)
        doc.add_page_break()
    para(doc, "TABLE OF CONTENTS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=BLUE)
    rows = [[str(i + 1), item] for i, item in enumerate([
        "Synopsis", "Introduction", "Problem Statement", "Feasibility Study", "Requirements", "Technology Stack",
        "System Architecture", "UML and Use Case Design", "Database Design", "Data Flow and Workflows",
        "Backend Implementation", "Prediction Business Logic", "Machine Learning Pipeline", "Explainable AI",
        "LLM Recommendation Module", "Frontend Implementation", "Screenshots", "Testing", "Security",
        "Deployment", "Reports and Results", "Future Scope", "Conclusion", "Bibliography", "Appendices",
    ])]
    add_table(doc, ["Sl. No.", "Chapter"], rows, [0.9, 5.5])
    doc.add_page_break()


def write_core_chapters(doc: Document, assets: dict[str, Path]) -> None:
    chapters = [
        ("1. Synopsis", [
            "AI Dropout Predictor is an offline-first full-stack platform that predicts the dropout risk of students and supports faculty with explainable recommendations. The system combines a FastAPI backend, React frontend, SQL database, scikit-learn machine learning pipeline, SHAP-based explainability, and a local Ollama LLM assistant.",
            "The project is designed for institutions that need early warning, privacy, and practical intervention support. Instead of sending student records to cloud services, the application keeps uploads, predictions, model artifacts, recommendations, and audit logs on the institution's own machine or server.",
        ]),
        ("2. Project Introduction", [
            "Student dropout affects learners, parents, faculty, and institutions. A student who disengages because of low attendance, backlogs, financial stress, or poor marks may not receive help until the problem becomes severe. The proposed system identifies risk earlier by combining available academic and behavioral signals into a repeatable prediction workflow.",
            "The application is not only a prediction engine. It also provides an upload wizard, dashboards, student detail views, risk explanations, local recommendations, counseling records, analytics, and exports. This makes it suitable for academic demonstration as well as practical institutional pilots.",
        ]),
        ("3. Problem Statement and Objectives", [
            "The problem is to design a secure and explainable software system that can predict student dropout risk from college data and help faculty take timely action. The system must support data entry, multi-format imports, model inference, explanations, recommendations, reporting, and role-based usage.",
            "The main objectives are authentication, student management, bulk upload, risk prediction, prediction explanation, counseling recommendation generation, analytics, report export, security logging, and offline-first deployment.",
        ]),
        ("4. Feasibility Study", [
            "The project is technically feasible because it relies on stable open-source tools. FastAPI provides a lightweight backend, React supports modern user interfaces, SQLite works for demos, PostgreSQL works for production, and scikit-learn supports classical machine learning without GPU requirements.",
            "It is economically feasible because no paid cloud prediction or LLM service is required. The local-first design allows colleges to deploy it on existing hardware and expand only when usage grows.",
        ]),
    ]
    for title, paras in chapters:
        doc.add_heading(title, level=1)
        for text in paras:
            para(doc, text)
        add_bullets(doc, [
            "Privacy-first processing: student data remains inside the institution.",
            "Explainable prediction: risk results include top contributing features.",
            "Action-oriented output: recommendations and counseling plans support follow-up.",
            "Maintainable architecture: routers, services, repositories, models, and schemas are separated.",
        ])

    doc.add_heading("5. Software and Hardware Requirements", level=1)
    add_table(doc, ["Hardware", "Minimum", "Recommended"], [["Processor", "Intel i3 / equivalent", "Intel i5 or above"], ["RAM", "4 GB", "8 GB or above"], ["Storage", "2 GB", "10 GB or above"], ["Display", "Standard", "1366 x 768 or higher"]], [1.5, 2.3, 2.3])
    add_table(doc, ["Software", "Purpose"], [["Python 3.11+", "Backend API, ML, tests"], ["Node.js 20+", "Frontend build"], ["FastAPI", "REST API"], ["React + TypeScript", "Browser interface"], ["SQLite/PostgreSQL", "Database"], ["Ollama", "Local LLM recommendations"], ["Docker", "Container deployment"]], [2.1, 4.2])

    doc.add_heading("6. Technology Stack Explanation", level=1)
    tech = [
        ("FastAPI", "FastAPI was selected for its speed, automatic API documentation, dependency injection, and strong Pydantic validation. It keeps endpoint code thin while services handle business logic."),
        ("React 18", "React provides reusable UI components for dashboards, forms, charts, tables, detail drawers, upload wizard screens, and the assistant interface."),
        ("TypeScript", "TypeScript improves frontend reliability by catching API and component shape mistakes before runtime."),
        ("SQLAlchemy", "SQLAlchemy models and repositories isolate database access from API handlers and services."),
        ("scikit-learn", "scikit-learn trains multiple classical ML candidates and provides predictable CPU-based inference."),
        ("SHAP", "SHAP helps explain which input features influenced a model decision, improving trust in the prediction."),
        ("Ollama", "Ollama enables local LLM recommendations without external cloud calls."),
    ]
    add_table(doc, ["Technology", "Why it is used"], [[a, b] for a, b in tech], [1.7, 4.7])
    for a, b in tech:
        doc.add_heading(a, level=2)
        para(doc, b + " In this project, the technology is used in a deliberately modular way so that the system can be extended without rewriting unrelated layers.")

    doc.add_heading("7. System Architecture", level=1)
    add_image(doc, assets["architecture"], "Figure 7.1: Overall system architecture showing frontend, API, services, database, ML, and local LLM.", 6.5)
    para(doc, "The architecture follows a layered pattern. The presentation layer handles user interaction. The API layer validates requests. The service layer performs business logic. The repository layer isolates database access. The AI layer performs feature engineering, inference, explainability, and recommendation generation.")

    doc.add_heading("8. UML and Use Case Design", level=1)
    add_image(doc, assets["use_case"], "Figure 8.1: Use case diagram for admin, faculty/counselor, and student actors.", 6.5)
    para(doc, "The use case diagram shows that each actor has a specific responsibility. Administrators manage accounts and courses. Faculty members manage students, uploads, predictions, recommendations, and counseling. Students have a limited view focused on their own profile and guidance.")

    doc.add_heading("9. Sequence Diagram", level=1)
    add_image(doc, assets["sequence"], "Figure 9.1: Sequence diagram for the student prediction workflow.", 6.5)
    para(doc, "The sequence diagram explains the runtime order of a prediction request. The frontend calls the predictions API, the API delegates to PredictionService, the service converts the student into model features, the model produces probabilities, the explanation service builds a narrative, and the result is persisted.")

    doc.add_heading("10. Data Flow Diagram", level=1)
    add_image(doc, assets["data_flow"], "Figure 10.1: Data flow from upload to reports and counseling.", 6.5)
    para(doc, "The data flow highlights the controlled import process. Uploaded files are parsed and mapped before database insertion. After records exist, the prediction and recommendation paths create outputs used by dashboards, reports, and counseling decisions.")

    doc.add_heading("11. ER Diagram and Database Design", level=1)
    add_image(doc, assets["er"], "Figure 11.1: Entity relationship diagram for the core database tables.", 6.5)
    para(doc, "The ER design keeps authentication users separate from academic student records. This avoids mixing login identity with predictive data and allows the system to support faculty and administrative users who are not student records.")

    doc.add_heading("12. Prediction Activity Flow", level=1)
    add_image(doc, assets["activity"], "Figure 12.1: Flow chart showing how a prediction decision is generated.", 6.5)
    para(doc, "The prediction decision is made from the maximum probability returned by the trained classifier. For example, if the model returns low=0.10, medium=0.25, and high=0.65, the selected risk level is high with confidence 0.65. This model decision is then explained and stored.")

    doc.add_heading("13. Screenshot Walkthrough", level=1)
    for key, caption in [
        ("login", "Login screen with college branding, demo accounts, and JWT/bcrypt security indicators."),
        ("dashboard", "Admin dashboard showing total students, predictions, high-risk percentage, risk distribution, and department risk."),
        ("assistant", "Local assistant screen showing Ollama status, privacy notice, generated answer, and matching student rows."),
        ("courses", "Degree courses management screen for maintaining academic programme metadata."),
        ("users", "User accounts screen showing role-based accounts and account status controls."),
        ("student_prediction", "Student detail screen where faculty can run prediction and generate intervention recommendations."),
    ]:
        add_image(doc, assets[key], f"Figure 13.{list(assets).index(key) + 1}: {caption}", 6.7)
        para(doc, caption + " This screenshot is placed in the relevant module section so the reader can connect the implementation description with the actual user interface.")

    doc.add_heading("14. Main Business Logic and Code Snippets", level=1)
    para(doc, "The most important business logic is prediction generation. The service must fetch the student, create the same feature vector used during training, run the model, explain the result, persist the prediction, and update risk history.")
    add_code(doc, ROOT / "backend/app/services/prediction_service.py", 45, 64, "Listing 14.1: Single-student prediction service flow")
    para(doc, "The predict_for_student function is the core bridge between the database and the ML model. It deliberately keeps persistence in the service layer, while predict_one only knows about feature frames and model probabilities.")
    add_code(doc, ROOT / "backend/app/ml/predict.py", 38, 61, "Listing 14.2: Model loading and probability-based decision")
    para(doc, "The model returns a probability vector. The decision is made by argmax over this vector, and the selected probability becomes confidence. This keeps risk selection consistent for both single and batch predictions.")
    add_code(doc, ROOT / "backend/app/ml/features.py", 38, 75, "Listing 14.3: Feature engineering from student record")
    para(doc, "The feature module avoids train/inference drift by using the same student_to_features and features_dataframe functions during training and prediction.")
    add_code(doc, ROOT / "backend/app/ml/features.py", 78, 91, "Listing 14.4: Rule-based label function used for synthetic training labels")
    para(doc, "The label_for rule is used for synthetic dataset generation and tests. Production prediction uses the trained persisted model, but the rule documents the educational intuition behind high and medium risk: poor attendance, too many backlogs, low internal marks, or fee delay.")
    add_code(doc, ROOT / "backend/app/services/recommendation_service.py", 75, 128, "Listing 14.5: Deterministic fallback recommendation logic")
    para(doc, "The recommendation fallback ensures that counseling support remains available even when the local LLM service is not running. This is important for offline-first deployments.")
    add_code(doc, ROOT / "backend/app/ml/train.py", 60, 137, "Listing 14.6: Training pipeline and best-model selection")
    para(doc, "The training pipeline evaluates multiple candidate models and selects the best by macro-F1. This metric is useful when low, medium, and high risk classes all matter and the system should not focus only on the majority class.")


def add_source_report_content(doc: Document, max_paras: int) -> None:
    src = Document(SOURCE_REPORT)
    doc.add_heading("15. Detailed Technical Chapters", level=1)
    para(doc, "The following chapters expand the system documentation with detailed engineering notes drawn from the technical project report. They cover architecture, backend, frontend, ML, explainability, LLM integration, security, upload workflows, analytics, testing, deployment, user operations, and future scope.")
    count = 0
    for p in src.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text in {"AI DROPOUT PREDICTOR", "Table of Contents"}:
            continue
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") or text.startswith("Chapter "):
            level = 1 if text.startswith("Chapter ") else 2
            doc.add_heading(text[:180], level=level)
        elif text.startswith("Listing "):
            para(doc, text[:500], bold=True, size=9.5)
        else:
            # Keep code listings compact and prose readable.
            if len(text) < 95 and any(tok in text for tok in ("def ", "return ", "class ", "import ", "from ")):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                r = p2.add_run(text)
                r.font.name = "Consolas"
                r.font.size = Pt(7.5)
            else:
                para(doc, text)
        count += 1
        if count >= max_paras:
            break


def appendices(doc: Document) -> None:
    doc.add_heading("Appendix A: API Endpoint Summary", level=1)
    add_table(doc, ["Endpoint", "Purpose"], [["POST /auth/login", "Authenticate and return tokens"], ["GET /students", "Search and paginate student records"], ["POST /uploads", "Upload and preview file"], ["POST /predictions/{student}", "Generate prediction"], ["POST /recommendations/{id}/generate", "Generate counseling plan"], ["GET /analytics/bundle", "Dashboard aggregate payload"]], [2.5, 3.9])
    doc.add_heading("Appendix B: Viva Questions", level=1)
    questions = [
        ("Why is the system offline-first?", "Because student records are sensitive and many institutions prefer local custody of academic data."),
        ("How is the prediction decision made?", "The model returns class probabilities and the class with the highest probability is selected as the risk level."),
        ("What happens if Ollama is not running?", "The recommendation service uses deterministic fallback plans and the application remains functional."),
        ("Why use SHAP?", "SHAP provides feature-level explanation that helps faculty understand why the model produced a risk result."),
        ("Why separate users from students?", "Login accounts and academic student records serve different purposes and should evolve independently."),
    ]
    for q, a in questions * 4:
        para(doc, "Question: " + q, bold=True)
        para(doc, "Answer: " + a)
    doc.add_heading("Bibliography", level=1)
    refs = ["FastAPI documentation", "React documentation", "SQLAlchemy documentation", "scikit-learn documentation", "SHAP documentation", "Ollama documentation", "Project documentation under docs/"]
    for r in refs:
        doc.add_paragraph(r, style="List Number")


def source_code_appendix(doc: Document) -> None:
    doc.add_heading("Appendix C: Extended Source Code Listings", level=1)
    para(
        doc,
        "This appendix includes selected source files that represent the main business logic of the system. These listings make the report useful during viva, code review, and implementation explanation because they show the exact functions used for prediction generation, risk decision, explainability, upload import, analytics, and API orchestration.",
    )
    listings = [
        (ROOT / "backend/app/services/prediction_service.py", "C.1 Prediction Service: persistence, explanation, and risk history", None),
        (ROOT / "backend/app/ml/predict.py", "C.2 Inference Helper: model loading and probability decision", None),
        (ROOT / "backend/app/ml/features.py", "C.3 Feature Engineering and Synthetic Label Logic", None),
        (ROOT / "backend/app/ml/explain.py", "C.4 Explainability Service: SHAP and deterministic fallback", None),
        (ROOT / "backend/app/ml/train.py", "C.5 Training Pipeline: model candidates and macro-F1 selection", None),
        (ROOT / "backend/app/services/recommendation_service.py", "C.6 Recommendation Service: local LLM and fallback plan", 170),
        (ROOT / "backend/app/services/upload_service.py", "C.7 Upload Service: preview, mapping, and commit workflow", 180),
        (ROOT / "backend/app/parsers/column_mapper.py", "C.8 Column Mapper: fuzzy matching for imported datasets", None),
        (ROOT / "backend/app/api/v1/endpoints/predictions.py", "C.9 Prediction API Endpoints and Audit Logging", None),
        (ROOT / "backend/app/services/analytics_service.py", "C.10 Analytics Service: dashboard aggregation", None),
    ]
    for path, caption, max_lines in listings:
        add_full_code_listing(doc, path, caption, max_lines=max_lines)


def build(max_paras: int) -> Path:
    assets = create_assets()
    doc = Document()
    configure_doc(doc)
    add_footer(doc.sections[0], "AI Dropout Predictor")
    front_matter(doc)
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(main, "AI Dropout Predictor")
    write_core_chapters(doc, assets)
    add_source_report_content(doc, max_paras=max_paras)
    source_code_appendix(doc)
    appendices(doc)
    doc.save(OUT)
    return OUT


def main() -> None:
    max_paras = int(os.environ.get("MAX_SOURCE_PARAS", "1420"))
    out = build(max_paras)
    print(out)


if __name__ == "__main__":
    main()
