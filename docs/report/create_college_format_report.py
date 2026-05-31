#!/usr/bin/env python3
"""Create a college-format project report for AI Dropout Predictor."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "AI_Dropout_Predictor_College_Format_Report.docx"
REPO = Path(__file__).resolve().parents[2]


BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number_footer(section, label: str) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(label + " | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    begin_run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = footer.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    instr_run._r.append(instr)

    end_run = footer.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 12, RGBColor(55, 65, 81), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def para(doc: Document, text: str = "", *, bold=False, align=None, size=None, color=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_center_title(doc: Document, text: str, size=16) -> None:
    p = para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, color=BLUE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = DARK
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_front_matter(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = para(doc, "AI DROPOUT PREDICTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=24, color=BLUE)
    p.paragraph_format.space_after = Pt(8)
    para(
        doc,
        "A Machine Learning Based System for Student Dropout Risk Prediction, Explainability, and Counseling Support",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=13,
        color=DARK,
    )
    doc.add_paragraph()
    para(doc, "PROJECT REPORT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=DARK)
    para(
        doc,
        "Submitted in partial fulfilment of the requirements for the Bachelor of Computer Applications / Computer Science programme",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        color=MUTED,
    )
    doc.add_paragraph()
    meta = [
        ["Submitted By", "[Student Name 1]   [Register No.]"],
        ["", "[Student Name 2]   [Register No.]"],
        ["", "[Student Name 3]   [Register No.]"],
        ["Guided By", "[Guide Name], Department of Computer Science"],
        ["Institution", "[College Name]"],
        ["Academic Year", f"{date.today():%Y}"],
    ]
    add_table(doc, ["Field", "Details"], meta, [1.8, 4.5])
    doc.add_page_break()

    add_center_title(doc, "DECLARATION")
    para(
        doc,
        "We hereby declare that the project titled \"AI Dropout Predictor\" has been developed by us in partial fulfilment of the academic requirements of the Computer Science programme. The work presented in this report is original to the best of our knowledge and has not been submitted previously for any degree or diploma to this or any other university.",
    )
    para(
        doc,
        "The system, documentation, and analysis have been completed under the guidance of the Department of Computer Science. All external references, technologies, and libraries used in the project have been acknowledged in the bibliography.",
    )
    doc.add_paragraph()
    add_table(doc, ["Place", "Date", "Student Signature"], [["[Place]", "[Date]", "[Name / Register No.]"], ["", "", "[Name / Register No.]"], ["", "", "[Name / Register No.]"]], [1.6, 1.4, 3.4])
    doc.add_page_break()

    add_center_title(doc, "PREFACE")
    para(
        doc,
        "The purpose of this project report is to present the design and implementation of AI Dropout Predictor, a full-stack application that helps educational institutions identify students who may be at risk of discontinuing their studies. The project combines data management, machine learning, explainable AI, and counseling support in a single offline-first platform.",
    )
    para(
        doc,
        "This report is written in a practical academic format. It explains the problem background, feasibility, requirements, system design, implementation modules, testing approach, results, and future enhancements. The content is based on the completed source code and documentation of the AI Dropout Predictor project.",
    )
    doc.add_page_break()

    add_center_title(doc, "ACKNOWLEDGEMENT")
    para(
        doc,
        "We express our sincere gratitude to the Department of Computer Science for providing the opportunity to complete this project. We are thankful to our project guide for valuable guidance, encouragement, and technical suggestions throughout the development process.",
    )
    para(
        doc,
        "We also thank our faculty members, friends, and family for their support and feedback. Their encouragement helped us complete the project with a focus on practical usefulness, privacy, and maintainable software design.",
    )
    doc.add_page_break()

    add_center_title(doc, "TABLE OF CONTENTS")
    rows = [
        ["1", "Synopsis"],
        ["2", "Project Introduction"],
        ["3", "Problem Statement and Objectives"],
        ["4", "Feasibility Study"],
        ["5", "Software and Hardware Requirements"],
        ["6", "Technology Overview"],
        ["7", "System Analysis and Design"],
        ["8", "Database Design"],
        ["9", "System Implementation"],
        ["10", "Machine Learning and Explainable AI"],
        ["11", "Testing"],
        ["12", "Results and Reports"],
        ["13", "Future Enhancements"],
        ["14", "Conclusion"],
        ["15", "Bibliography"],
    ]
    add_table(doc, ["Sl. No.", "Contents"], rows, [1.0, 5.2])
    doc.add_page_break()


def add_project_content(doc: Document) -> None:
    doc.add_heading("1. Synopsis", level=1)
    para(
        doc,
        "AI Dropout Predictor is an offline-first web application designed to predict student dropout risk and assist faculty members with timely interventions. The platform stores student academic and behavioural data, runs machine learning predictions, explains the reasons behind each risk score, and produces counseling recommendations using a local LLM with deterministic fallback logic.",
    )
    add_table(
        doc,
        ["Item", "Description"],
        [
            ["Project Title", "AI Dropout Predictor"],
            ["Project Type", "Full-stack machine learning web application"],
            ["Primary Users", "Administrator, faculty/counselor, student"],
            ["Core Purpose", "Identify at-risk students early and support counseling decisions"],
            ["Deployment Mode", "Local/offline-first; SQLite by default, PostgreSQL optional"],
            ["Main Technologies", "FastAPI, React, TypeScript, SQLAlchemy, scikit-learn, SHAP, Ollama"],
        ],
        [1.8, 4.6],
    )
    para(
        doc,
        "The application is suitable for colleges that want an early-warning system without sending student records to third-party cloud services. It supports CSV, Excel, PDF, and DOCX data uploads, maps columns using fuzzy matching, persists predictions, and displays analytics dashboards for departments and individual students.",
    )

    doc.add_heading("2. Project Introduction", level=1)
    para(
        doc,
        "Student dropout is a serious concern for higher education institutions because it affects student careers, institutional performance, and academic planning. Traditional monitoring methods depend on manual review of attendance, marks, and fee records. Such methods are slow, inconsistent, and often detect problems only after the student has already disengaged.",
    )
    para(
        doc,
        "AI Dropout Predictor solves this problem by providing a single platform where institutions can manage student records, calculate dropout risk, view explainability factors, and generate intervention plans. The system focuses on practical signals such as attendance percentage, internal marks, semester marks, backlogs, fee delay, financial background, placement readiness, and engagement score.",
    )
    doc.add_heading("2.1 Need for the System", level=2)
    add_bullets(
        doc,
        [
            "Manual tracking becomes difficult when the number of students increases.",
            "Faculty need early and explainable alerts instead of late semester summaries.",
            "Student records are sensitive and should remain within the institution.",
            "Counseling decisions need evidence, consistency, and follow-up tracking.",
            "Administrative users need department-level analytics and exportable reports.",
        ],
    )
    doc.add_heading("2.2 Existing System", level=2)
    para(
        doc,
        "In many colleges, attendance registers, marksheets, and fee records are maintained separately. Faculty members inspect these records manually or use spreadsheets for basic filtering. This approach does not provide predictive risk scoring, automated explanations, or structured counseling recommendations.",
    )
    doc.add_heading("2.3 Proposed System", level=2)
    para(
        doc,
        "The proposed system integrates data ingestion, prediction, explainability, recommendations, counseling logs, analytics, and reporting. It runs on institutional hardware and uses local model artifacts. LLM-based features communicate with a local Ollama service, and when the LLM is unavailable, the application continues with rule-based fallback recommendations.",
    )

    doc.add_heading("3. Problem Statement and Objectives", level=1)
    para(
        doc,
        "The problem addressed by this project is the lack of an accessible, privacy-preserving, and explainable early-warning system for student dropout risk. The system must predict risk from available institutional data and help faculty act before the risk becomes irreversible.",
    )
    doc.add_heading("3.1 Objectives", level=2)
    add_numbered(
        doc,
        [
            "Develop a secure login system with role-based access for admin, faculty, and student users.",
            "Allow student records to be created manually or uploaded from common file formats.",
            "Train and serve a machine learning model for dropout risk classification.",
            "Explain every prediction using SHAP or a deterministic fallback method.",
            "Generate actionable counseling plans through a local LLM or offline fallback.",
            "Provide analytics dashboards and exportable PDF/Excel reports.",
            "Maintain audit logs and deployment scripts for practical institutional use.",
        ],
    )

    doc.add_heading("4. Feasibility Study", level=1)
    add_table(
        doc,
        ["Feasibility Area", "Assessment"],
        [
            ["Technical Feasibility", "The project uses mature open-source tools: FastAPI, React, SQLAlchemy, scikit-learn, and SQLite/PostgreSQL. These technologies are stable, documented, and suitable for local deployment."],
            ["Operational Feasibility", "Faculty can upload familiar spreadsheet files, review mapped fields, run predictions, and view recommendations through a browser-based interface."],
            ["Economic Feasibility", "The application can run on a normal laptop or institutional server. It does not require paid cloud APIs for prediction or recommendations."],
            ["Schedule Feasibility", "The system is modular, allowing backend, frontend, ML, and documentation work to proceed independently."],
            ["Security Feasibility", "JWT authentication, bcrypt password hashing, rate limiting, role gates, and audit logs provide a secure base for academic deployment."],
        ],
        [1.7, 4.7],
    )

    doc.add_heading("5. Software and Hardware Requirements", level=1)
    doc.add_heading("5.1 Hardware Requirements", level=2)
    add_table(
        doc,
        ["Component", "Minimum Requirement", "Recommended"],
        [
            ["Processor", "Intel i3 / equivalent", "Intel i5 or above"],
            ["RAM", "4 GB", "8 GB or above"],
            ["Storage", "2 GB free space", "10 GB free space for logs/uploads"],
            ["Display", "Standard monitor", "1366 x 768 or higher"],
            ["Network", "Localhost for demo", "LAN/HTTPS for institution deployment"],
        ],
        [1.6, 2.3, 2.3],
    )
    doc.add_heading("5.2 Software Requirements", level=2)
    add_table(
        doc,
        ["Software", "Use"],
        [
            ["Windows 10/11 or Linux", "Operating system for development and deployment"],
            ["Python 3.11+", "Backend API, ML, parsers, and tests"],
            ["Node.js 20+", "Frontend development and build tooling"],
            ["FastAPI", "REST API framework"],
            ["React 18 + Vite", "Single-page web application"],
            ["SQLite / PostgreSQL", "Database storage"],
            ["Ollama", "Optional local LLM runtime for recommendations and chat"],
            ["Docker", "Containerized deployment"],
        ],
        [2.0, 4.3],
    )

    doc.add_heading("6. Technology Overview", level=1)
    para(
        doc,
        "FastAPI is used for the backend because it supports typed request and response models through Pydantic and automatically generates OpenAPI documentation. React and TypeScript provide a maintainable frontend with reusable components. SQLAlchemy separates database operations from business logic. scikit-learn provides model training and inference, while SHAP improves transparency by identifying feature contributions.",
    )
    add_table(
        doc,
        ["Layer", "Technologies", "Purpose"],
        [
            ["Frontend", "React, TypeScript, Vite, Tailwind, React Query", "Browser interface, dashboards, forms, charts"],
            ["Backend", "FastAPI, Pydantic, SQLAlchemy", "API, validation, business logic, persistence"],
            ["Machine Learning", "scikit-learn, optional XGBoost, joblib", "Training, model selection, prediction"],
            ["Explainability", "SHAP, permutation fallback", "Human-readable risk reasons"],
            ["LLM", "Ollama with fallback templates", "Counseling recommendations and chat support"],
            ["Deployment", "Docker, PowerShell/Bash scripts", "Repeatable setup and startup"],
        ],
        [1.4, 2.5, 2.5],
    )

    doc.add_heading("7. System Analysis and Design", level=1)
    doc.add_heading("7.1 User Roles", level=2)
    add_table(
        doc,
        ["Actor", "Responsibilities"],
        [
            ["Administrator", "Manage users, departments, model retraining, settings, reports, and global analytics."],
            ["Faculty / Counselor", "Manage department students, upload data, run predictions, review explanations, generate recommendations, and record counseling sessions."],
            ["Student", "View personal profile, prediction history, and recommendations where enabled."],
        ],
        [1.7, 4.7],
    )
    doc.add_heading("7.2 Functional Requirements", level=2)
    add_bullets(
        doc,
        [
            "The system shall authenticate users and issue access and refresh tokens.",
            "The system shall support CRUD operations for students and users.",
            "The system shall preview and map uploaded files before inserting records.",
            "The system shall run single-student and batch predictions.",
            "The system shall store predictions, risk history, explanations, and recommendations.",
            "The system shall provide analytics and export reports.",
            "The system shall log sensitive actions for audit purposes.",
        ],
    )
    doc.add_heading("7.3 Use Case Summary", level=2)
    add_table(
        doc,
        ["Use Case", "Primary Actor", "Description"],
        [
            ["Login", "All users", "User enters email and password and receives a secure session."],
            ["Upload Student Data", "Faculty/Admin", "Upload CSV, Excel, PDF, or DOCX and confirm mapped columns."],
            ["Run Prediction", "Faculty/Admin", "Generate dropout risk and explanation for a student."],
            ["Generate Recommendation", "Faculty/Admin", "Create a counseling plan from prediction context."],
            ["View Analytics", "Faculty/Admin", "Review risk distribution and department trends."],
            ["Export Report", "Faculty/Admin", "Download PDF or Excel reports."],
        ],
        [1.7, 1.5, 3.2],
    )
    doc.add_heading("7.4 Architecture", level=2)
    para(
        doc,
        "The system follows a three-tier architecture. The React frontend communicates with the FastAPI backend using JSON REST APIs. The backend uses SQLAlchemy for database access, calls ML services for prediction, and optionally communicates with local Ollama for recommendations and chat.",
    )
    add_table(
        doc,
        ["Tier", "Components"],
        [
            ["Presentation Tier", "React pages, components, charts, upload wizard, role-aware navigation"],
            ["Application Tier", "FastAPI routers, services, repositories, schemas, authentication, rate limits"],
            ["Data and AI Tier", "SQLite/PostgreSQL database, ML artifacts, SHAP explainer, Ollama local LLM"],
        ],
        [1.6, 4.8],
    )

    doc.add_heading("8. Database Design", level=1)
    para(
        doc,
        "The database design centers on students, departments, predictions, recommendations, counseling sessions, uploads, users, and audit logs. Predictions store the risk level, confidence, model version, features JSON, and explanation JSON. Risk history supports trend charts, and audit logs provide traceability for sensitive operations.",
    )
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["users", "Stores login identity, role, password hash, and active status."],
            ["departments", "Groups students and faculty by academic department."],
            ["students", "Stores academic, demographic, attendance, and fee-related student fields."],
            ["uploads", "Tracks uploaded files, mapping status, and import metadata."],
            ["predictions", "Stores dropout risk result, confidence, feature payload, and explanation."],
            ["risk_history", "Maintains time-series risk snapshots for charts."],
            ["recommendations", "Stores LLM or fallback counseling recommendations."],
            ["counseling_sessions", "Records counseling notes, follow-up dates, and status."],
            ["audit_logs", "Records important user and system actions."],
        ],
        [1.8, 4.6],
    )

    doc.add_heading("9. System Implementation", level=1)
    doc.add_heading("9.1 Backend Modules", level=2)
    add_bullets(
        doc,
        [
            "API endpoints validate requests and delegate work to services.",
            "Service classes contain business logic for students, uploads, predictions, recommendations, reports, analytics, and chat.",
            "Repository classes centralize database queries.",
            "Pydantic schemas define request and response contracts.",
            "Core modules handle configuration, security, dependencies, rate limiting, and logging.",
        ],
    )
    doc.add_heading("9.2 Frontend Modules", level=2)
    add_bullets(
        doc,
        [
            "App shell provides sidebar navigation, topbar, theme control, and protected routes.",
            "Students page supports search, filtering, forms, and a detail drawer.",
            "Uploads page implements a three-step import wizard with column mapping.",
            "Predictions page displays risk meter, explainability panel, and recommendations.",
            "Analytics page uses charts for risk distribution, department risk, attendance trends, and feature importance.",
            "Chat page provides local assistant interaction with offline status handling.",
        ],
    )
    doc.add_heading("9.3 File Upload and Mapping", level=2)
    para(
        doc,
        "The upload workflow validates file type and MIME information, extracts tabular rows using the appropriate parser, and suggests target fields using fuzzy matching. Users can review the preview before committing records to the database, reducing the risk of accidental bad imports.",
    )

    doc.add_heading("10. Machine Learning and Explainable AI", level=1)
    para(
        doc,
        "The ML pipeline converts student attributes into numerical and categorical features, trains multiple candidate classifiers, selects the best model by macro-F1 score, and persists the chosen artifact with metadata. At prediction time, the model returns a risk level and confidence score.",
    )
    add_table(
        doc,
        ["Feature", "Meaning"],
        [
            ["attendance_pct", "Attendance percentage"],
            ["internal_marks", "Internal assessment marks"],
            ["semester_marks", "Semester or final marks"],
            ["backlogs", "Number of uncleared subjects"],
            ["fee_delay_days", "Delay in fee payment"],
            ["fee_paid", "Whether fee is paid"],
            ["financial_status", "Financial background indicator"],
            ["placement_readiness", "Career readiness category"],
            ["engagement_score", "Participation and engagement indicator"],
        ],
        [2.1, 4.3],
    )
    para(
        doc,
        "Explainability is an important part of this project because faculty should not see only a black-box label. The system returns top contributing factors and a plain-language explanation so the counselor can understand whether the risk is mainly due to attendance, marks, backlogs, fee delay, or engagement.",
    )

    doc.add_heading("11. Testing", level=1)
    para(
        doc,
        "Testing was performed at API, service, and workflow levels. The backend includes pytest tests for health checks, authentication, student CRUD, upload preview, and prediction flow. Manual testing was also performed from the frontend to verify login, navigation, data upload, prediction, recommendations, analytics, and report export.",
    )
    add_table(
        doc,
        ["Test Case", "Expected Result", "Status"],
        [
            ["Valid user login", "Access and refresh tokens are returned", "Pass"],
            ["Invalid login", "Request is rejected with an error", "Pass"],
            ["Student creation", "Student record is saved and visible in list", "Pass"],
            ["CSV upload preview", "Rows and suggested mappings are displayed", "Pass"],
            ["Prediction request", "Risk level, confidence, and explanation are returned", "Pass"],
            ["LLM unavailable", "Fallback recommendation is generated", "Pass"],
            ["Role restriction", "Unauthorized pages/actions are blocked", "Pass"],
            ["Report export", "PDF/Excel report is generated", "Pass"],
        ],
        [2.2, 3.0, 1.0],
    )

    doc.add_heading("12. Results and Reports", level=1)
    para(
        doc,
        "The completed system provides a working demonstration of an institutional dropout-risk platform. Users can add students, upload bulk records, run predictions, inspect explanations, create recommendations, and view charts. Reports can be exported for students and departments, helping faculty discuss interventions in meetings or counseling sessions.",
    )
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["Risk Meter", "Shows low, medium, or high dropout risk with confidence."],
            ["Explainability Panel", "Lists top factors that increased or decreased risk."],
            ["Recommendation Plan", "Provides suggested actions, priority, and follow-up guidance."],
            ["Analytics Dashboard", "Displays aggregate risk distribution and trends."],
            ["PDF/Excel Reports", "Allows data to be shared outside the application."],
        ],
        [1.8, 4.6],
    )

    doc.add_heading("13. Future Enhancements", level=1)
    add_bullets(
        doc,
        [
            "Integrate with college ERP systems for automatic attendance and marks synchronization.",
            "Add IoT or biometric attendance feeds for real-time risk updates.",
            "Develop a mobile app for faculty and students.",
            "Add parent/guardian notification workflows with consent controls.",
            "Introduce continuous learning from counseling outcomes.",
            "Add richer multilingual LLM support for local counseling contexts.",
            "Generate official academic reports with institution branding.",
        ],
    )

    doc.add_heading("14. Conclusion", level=1)
    para(
        doc,
        "AI Dropout Predictor demonstrates how machine learning and explainable AI can be applied to a real educational problem in a privacy-preserving way. The project does not simply classify students; it also explains why a student may be at risk and supports faculty action through recommendations and counseling records.",
    )
    para(
        doc,
        "The system is modular, offline-first, and practical for academic institutions. Its combination of FastAPI, React, SQLAlchemy, scikit-learn, SHAP, and local LLM support makes it a strong foundation for future expansion into ERP integration, mobile access, and continuous intervention tracking.",
    )

    doc.add_heading("15. Bibliography", level=1)
    refs = [
        "FastAPI documentation, https://fastapi.tiangolo.com/",
        "React documentation, https://react.dev/",
        "SQLAlchemy documentation, https://www.sqlalchemy.org/",
        "scikit-learn documentation, https://scikit-learn.org/",
        "SHAP documentation, https://shap.readthedocs.io/",
        "Ollama documentation, https://ollama.com/",
        "Project repository documentation: docs/architecture.md, docs/api.md, docs/ml-pipeline.md, docs/security.md.",
    ]
    add_numbered(doc, refs)


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A: Project Folder Structure", level=1)
    entries = [
        ["backend/app/api/v1/endpoints", "REST API route modules"],
        ["backend/app/services", "Business logic layer"],
        ["backend/app/repositories", "Database query layer"],
        ["backend/app/ml", "Feature engineering, training, prediction, explainability"],
        ["backend/app/parsers", "CSV, Excel, PDF, DOCX parsing and column mapping"],
        ["frontend/src/pages", "Main routed frontend pages"],
        ["frontend/src/components", "Reusable UI, charts, layout, and feature components"],
        ["docs", "Architecture, API, security, ML, LLM, deployment, and future scope documentation"],
        ["scripts", "Setup, startup, model training, and LLM helper scripts"],
    ]
    add_table(doc, ["Path", "Purpose"], entries, [2.6, 3.8])

    doc.add_heading("Appendix B: Sample Login Credentials", level=1)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ["Admin", "admin@example.com", "Admin@123"],
            ["Faculty", "faculty@example.com", "Faculty@123"],
            ["Student", "student@example.com", "Student@123"],
        ],
        [1.4, 2.6, 2.2],
    )
    para(doc, "Note: These demo credentials should be changed before any real deployment.")


def main() -> None:
    doc = Document()
    configure_doc(doc)
    add_page_number_footer(doc.sections[0], "AI Dropout Predictor")
    add_front_matter(doc)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_page_number_footer(main_section, "AI Dropout Predictor")
    add_project_content(doc)
    add_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
